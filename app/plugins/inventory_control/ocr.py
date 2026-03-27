from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any

try:
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover - optional dependency at runtime
    pytesseract = None
    Image = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None  # type: ignore

try:
    import fitz  # PyMuPDF - optional, for scanned PDFs (render page to image + OCR)
except Exception:
    fitz = None  # type: ignore

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None  # type: ignore


@dataclass
class ParsedInvoiceLine:
    sku: str
    description: str
    quantity: float
    unit_price: float
    line_total: float
    external_item_ref: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedInvoice:
    source: str
    invoice_number: Optional[str]
    invoice_date: Optional[str]
    currency: Optional[str]
    supplier_name: Optional[str]
    lines: List[ParsedInvoiceLine]
    raw_text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class OCRProvider:
    def extract_text(self, file_path: str) -> str:
        raise NotImplementedError


class TesseractOCRProvider(OCRProvider):
    """
    Thin wrapper around pytesseract so we can swap OCR providers later.
    """

    def extract_text(self, file_path: str) -> str:
        if pytesseract is None or Image is None:
            raise RuntimeError(
                "pytesseract/Pillow not available; install them to enable OCR."
            )
        if not os.path.exists(file_path):
            raise FileNotFoundError(file_path)
        img = Image.open(file_path)
        try:
            return pytesseract.image_to_string(img) or ""
        finally:
            img.close()


# ---------------------------------------------------------------------------
# PDF / DOCX text extraction (in addition to image OCR)
# ---------------------------------------------------------------------------

def _pdf_extract_text_pypdf(file_path: str) -> str:
    """Extract text from PDF using pypdf (text-based PDFs)."""
    if PdfReader is None:
        raise RuntimeError("pypdf not available; install pypdf for PDF support.")
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _pdf_extract_text_ocr_fallback(file_path: str, tesseract_provider: TesseractOCRProvider) -> str:
    """Render each PDF page to an image and run OCR (for scanned PDFs). Requires PyMuPDF and pytesseract."""
    if fitz is None:
        raise RuntimeError("PyMuPDF (pymupdf) not available; install pymupdf for scanned PDF support.")
    if pytesseract is None or Image is None:
        raise RuntimeError("pytesseract/Pillow required for PDF OCR fallback.")
    doc = fitz.open(file_path)
    parts = []
    try:
        for i in range(len(doc)):
            page = doc[i]
            pix = page.get_pixmap(dpi=150, alpha=False)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                pix.save(tmp_path)
                text = tesseract_provider.extract_text(tmp_path)
                if text:
                    parts.append(text)
            finally:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
    finally:
        doc.close()
    return "\n".join(parts)


def _docx_extract_text(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    if DocxDocument is None:
        raise RuntimeError("python-docx not available; install python-docx for DOCX support.")
    doc = DocxDocument(file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_file(
    file_path: str,
    ocr_provider: Optional[OCRProvider] = None,
) -> str:
    """
    Extract text from a file supporting images (OCR), PDF, and DOCX.
    - Images (.png, .jpg, .jpeg, .gif, .tiff, .bmp): OCR via ocr_provider (default Tesseract).
    - PDF: pypdf text extraction; if very little text (scanned PDF), fall back to PyMuPDF render + OCR if available.
    - DOCX: python-docx (paragraphs and table cells).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(file_path)
    ext = (os.path.splitext(file_path)[1] or "").lower()
    provider = ocr_provider or TesseractOCRProvider()

    if ext == ".pdf":
        # Prefer pypdf (no extra deps). If almost no text, try scanned-PDF path.
        try:
            text = _pdf_extract_text_pypdf(file_path)
        except Exception:
            text = ""
        if len((text or "").strip()) < 80 and fitz is not None:
            try:
                text = _pdf_extract_text_ocr_fallback(
                    file_path,
                    tesseract_provider=provider if isinstance(provider, TesseractOCRProvider) else TesseractOCRProvider(),
                )
            except Exception:
                pass
        return text or ""

    if ext in (".docx", ".doc"):
        # .doc (old binary) not supported by python-docx; only .docx
        if ext == ".doc":
            raise ValueError("Binary .doc format is not supported; use .docx or export to PDF.")
        return _docx_extract_text(file_path)

    # Image types: use OCR
    return provider.extract_text(file_path)


class InvoiceParser:
    def parse(self, raw_text: str, source: str) -> ParsedInvoice:
        raise NotImplementedError


class AmazonInvoiceParser(InvoiceParser):
    """
    Very lightweight Amazon-like invoice parser.

    This is intentionally conservative: it focuses on extracting line items
    where quantity, description, and price appear in a simple delimited form.
    Real-world production deployments can replace or extend this with a more
    robust parser while keeping the same ParsedInvoice contract.
    """

    def parse(self, raw_text: str, source: str = "amazon") -> ParsedInvoice:
        lines = []
        invoice_number = None
        invoice_date = None
        currency = None
        supplier_name = "Amazon"

        for line in raw_text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue

            lower = stripped.lower()
            if "invoice" in lower and "number" in lower and ":" in stripped:
                # e.g. "Invoice Number: 123-1234567-1234567"
                invoice_number = stripped.split(":", 1)[1].strip()
                continue
            if "order date" in lower or "invoice date" in lower:
                # naive capture of date portion after colon
                if ":" in stripped:
                    invoice_date = stripped.split(":", 1)[1].strip()
                continue

            # Very simple line item heuristic:
            # "SKU123 | Widget Name | Qty 2 | £10.00"
            parts = [p.strip() for p in stripped.split("|")]
            if len(parts) >= 3 and "qty" in parts[2].lower():
                sku = parts[0]
                description = parts[1]
                qty_part = parts[2].lower().replace("qty", "").strip()
                try:
                    quantity = float(qty_part.split()[0])
                except Exception:
                    continue

                unit_price = 0.0
                line_total = 0.0
                if len(parts) >= 4:
                    price_str = parts[3].replace("£", "").replace("$", "").strip()
                    try:
                        unit_price = float(price_str)
                    except Exception:
                        unit_price = 0.0

                if unit_price and quantity:
                    line_total = unit_price * quantity

                lines.append(
                    ParsedInvoiceLine(
                        sku=sku,
                        description=description,
                        quantity=quantity,
                        unit_price=unit_price,
                        line_total=line_total,
                    )
                )

        return ParsedInvoice(
            source=source,
            invoice_number=invoice_number,
            invoice_date=invoice_date,
            currency=currency,
            supplier_name=supplier_name,
            lines=lines,
            raw_text=raw_text,
            metadata={},
        )


class InventoryInvoiceService:
    """
    Service that coordinates OCR + parsing and persistence.

    Persistence into inventory_invoices and inventory_invoice_lines is added
    later together with the API layer; for now this class focuses on the
    OCR/parse pipeline so it can be exercised in tests.
    """

    def __init__(self, ocr_provider: OCRProvider | None = None, parser: InvoiceParser | None = None):
        self.ocr_provider = ocr_provider or TesseractOCRProvider()
        self.parser = parser or AmazonInvoiceParser()

    def parse_file(self, file_path: str, source: str = "amazon") -> ParsedInvoice:
        """
        High-level helper used by routes and tests:
        - Extracts text from the file (images via OCR, PDF via pypdf/OCR fallback, DOCX via python-docx)
        - Parses the resulting text into a structured ParsedInvoice
        """
        text = extract_text_from_file(file_path, ocr_provider=self.ocr_provider)
        return self.parser.parse(text, source=source)

